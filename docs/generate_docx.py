from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# --- Page margins ---
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.54)

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    return h

def add_para(text, bold=False, italic=False, size=None, alignment=None, space_after=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if alignment:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '1A1A2E')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r, row_data in enumerate(rows):
        for c, value in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            run = cell.paragraphs[0].add_run(str(value))
            run.font.size = Pt(9)
            if r % 2 == 0:
                set_cell_shading(cell, 'F0F4FF')

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, width in enumerate(col_widths):
                row.cells[i].width = Cm(width)

    doc.add_paragraph()  # spacing after table
    return table

# ==========================================
# DOCUMENT CONTENT
# ==========================================

# --- Title ---
add_heading_styled('Configuración del entorno virtual', level=1)

# --- Description ---
add_heading_styled('Descripción', level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run = p.add_run(
    'Se estructurará un entorno virtual basado en una arquitectura de microservicios, '
    'desplegado mediante contenedores Docker para garantizar escalabilidad, portabilidad y '
    'aislamiento entre componentes. El diseño incluye un API Gateway como punto de entrada '
    'único con validación JWT y resolución de tenant por subdominio, ocho servicios independientes '
    '(API Gateway, Auth, Core, Booking, Payment, Notification, Marketplace y Analytics), bases de '
    'datos PostgreSQL separadas por servicio, y servicios complementarios de caché (Redis) y '
    'mensajería asíncrona (RabbitMQ).'
)
run.font.size = Pt(11)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
run2 = p2.add_run(
    'Se implementará un firewall perimetral con pfSense y segmentación de red en VLANs para '
    'controlar accesos, aislar las capas de aplicación, datos y auditoría, y proteger los '
    'componentes críticos, asegurando que posibles vulnerabilidades no comprometan la '
    'infraestructura principal ni los datos multi-tenant de los negocios alojados en la plataforma.'
)
run2.font.size = Pt(11)

# --- Topology ---
add_heading_styled('Topología de red', level=2)

topo = (
    '                              INTERNET\n'
    '                                 |\n'
    '                     +-----------v----------+\n'
    '                     |     Gateway-FW       |\n'
    '                     |     pfSense 2.7      |\n'
    '                     |    192.168.30.1       |\n'
    '                     |  Firewall / NAT / RP  |\n'
    '                     +-----------+----------+\n'
    '                                 |\n'
    '          +----------------------+----------------------+\n'
    '          |                      |                      |\n'
    ' +--------v---------+  +--------v---------+  +--------v---------+\n'
    ' |    VLAN 20       |  |    VLAN 30       |  |    VLAN 40       |\n'
    ' |   Aplicación     |  |     Datos        |  |   Auditoría      |\n'
    ' |  10.0.20.0/24   |  |  10.0.30.0/24   |  |  10.0.40.0/24   |\n'
    ' +--------+---------+  +--------+---------+  +--------+---------+\n'
    '          |                      |                      |\n'
    ' +--------v---------+  +--------v---------+  +--------v---------+\n'
    ' | App-Srv-01       |  |  DB-Server       |  |  Kali-Auditor    |\n'
    ' | 10.0.20.10       |  |  10.0.30.10      |  |  10.0.40.10      |\n'
    ' +------------------+  +------------------+  +------------------+\n'
    ' +--------v---------+  +--------v---------+\n'
    ' | App-Srv-02       |  |  Redis-Cache     |\n'
    ' | 10.0.20.11       |  |  10.0.30.11      |\n'
    ' +------------------+  +------------------+\n'
    ' +--------v---------+  +--------v---------+\n'
    ' | App-Srv-03       |  |  RabbitMQ-Broker |\n'
    ' | 10.0.20.12       |  |  10.0.30.12      |\n'
    ' +------------------+  +------------------+'
)

topo_para = doc.add_paragraph()
topo_run = topo_para.add_run(topo)
topo_run.font.name = 'Consolas'
topo_run.font.size = Pt(7.5)
topo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_para('Figura 1. Topología de red del entorno virtual BeautySpot', italic=True, size=9,
         alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

# --- VM Table ---
add_heading_styled('Distribución de máquinas virtuales', level=2)

vm_headers = ['Nombre de la VM', 'Sistema Operativo', 'Dirección IP', 'VLAN', 'Rol y servicios']
vm_rows = [
    ['Gateway-FW', 'pfSense 2.7', '192.168.30.1', 'WAN',
     'Firewall perimetral, NAT, reverse proxy con resolución de subdominios ({slug}.beautyspot.local), certificados SSL/TLS y rate limiting de entrada'],
    ['App-Srv-01', 'Ubuntu Server 22.04 LTS', '10.0.20.10', '20 - Aplicación',
     'Servidor de aplicación (Docker): API Gateway (3000), Auth Service (3001), Core Service (3002) + Frontend Next.js (8080)'],
    ['App-Srv-02', 'Ubuntu Server 22.04 LTS', '10.0.20.11', '20 - Aplicación',
     'Servidor de aplicación (Docker): Booking Service (3003), Payment Service (3004), Notification Service (3005)'],
    ['App-Srv-03', 'Ubuntu Server 22.04 LTS', '10.0.20.12', '20 - Aplicación',
     'Servidor de aplicación (Docker): Marketplace Service (3006), Analytics Service (3007)'],
    ['DB-Server', 'Ubuntu Server 22.04 LTS', '10.0.30.10', '30 - Datos',
     'Servidor de base de datos: PostgreSQL 16 con 7 bases de datos independientes (beautyspot_auth, beautyspot_core, beautyspot_booking, beautyspot_payment, beautyspot_notification, beautyspot_marketplace, beautyspot_analytics)'],
    ['Redis-Cache', 'Ubuntu Server 22.04 LTS', '10.0.30.11', '30 - Datos',
     'Servidor de caché y sesiones: Redis 7 con persistencia AOF para resolución de tenant, sesiones JWT y datos de disponibilidad en tiempo real'],
    ['RabbitMQ-Broker', 'Ubuntu Server 22.04 LTS', '10.0.30.12', '30 - Datos',
     'Broker de mensajería: RabbitMQ 3 con management plugin para eventos de dominio entre microservicios (appointment.created, payment.registered, etc.)'],
    ['Kali-Auditor', 'Kali Linux 2024.1', '10.0.40.10', '40 - Auditoría',
     'Equipo de auditoría para pruebas de seguridad: análisis de vulnerabilidades (Nmap, Nikto), pruebas de penetración sobre la API, validación de JWT, verificación de aislamiento multi-tenant y pruebas de inyección SQL/XSS'],
    ['Metasploit-Edu', 'Metasploitable 3', '10.0.40.11', '40 - Auditoría',
     'Servidor de pruebas controlado: objetivo de práctica para validar que las medidas de seguridad del firewall, segmentación de red y hardening de contenedores Docker funcionan correctamente'],
]
add_table(vm_headers, vm_rows, col_widths=[2.8, 2.5, 2.0, 2.0, 7.5])

# --- Firewall Rules ---
add_heading_styled('Reglas de firewall (pfSense)', level=2)

fw_headers = ['#', 'Origen', 'Destino', 'Puertos', 'Protocolo', 'Descripción']
fw_rows = [
    ['1', 'WAN (Internet)', 'Gateway-FW', '80, 443', 'TCP', 'Tráfico HTTP/HTTPS entrante al reverse proxy'],
    ['2', 'Gateway-FW', 'VLAN 20 (App)', '3000-3007, 8080', 'TCP', 'Enrutamiento a microservicios y frontend'],
    ['3', 'VLAN 20 (App)', 'VLAN 30 (Datos)', '5432, 6379, 5672, 15672', 'TCP', 'Acceso de aplicación a PostgreSQL, Redis y RabbitMQ'],
    ['4', 'VLAN 20 (App)', 'VLAN 20 (App)', '3000-3007', 'TCP', 'Comunicación interna entre microservicios'],
    ['5', 'VLAN 40 (Auditoría)', 'VLAN 20 (App)', '3000-3007, 8080', 'TCP', 'Escaneo de puertos y pruebas sobre servicios de aplicación'],
    ['6', 'VLAN 40 (Auditoría)', 'VLAN 30 (Datos)', '5432', 'TCP', 'Auditoría directa de configuración de PostgreSQL'],
    ['7', 'VLAN 30 (Datos)', '*', '*', '*', 'DENEGAR — Los servidores de datos no inician conexiones hacia otras VLANs'],
    ['8', 'VLAN 20 (App)', 'WAN', '80, 443', 'TCP', 'Salida controlada para descarga de dependencias (solo durante despliegue)'],
]
add_table(fw_headers, fw_rows, col_widths=[0.7, 2.5, 2.5, 2.8, 1.5, 6.8])

# --- Resources ---
add_heading_styled('Recursos asignados por VM', level=2)

res_headers = ['VM', 'CPU (vCPU)', 'RAM (GB)', 'Disco (GB)', 'Justificación']
res_rows = [
    ['Gateway-FW', '1', '1', '16', 'Firewall ligero con pfSense'],
    ['App-Srv-01', '2', '4', '40', '3 microservicios NestJS + frontend Next.js'],
    ['App-Srv-02', '2', '4', '30', '3 microservicios NestJS'],
    ['App-Srv-03', '1', '2', '30', '2 microservicios NestJS'],
    ['DB-Server', '2', '4', '60', '7 bases de datos PostgreSQL con índices y datos multi-tenant'],
    ['Redis-Cache', '1', '2', '16', 'Caché en memoria con persistencia AOF'],
    ['RabbitMQ-Broker', '1', '2', '16', 'Broker de mensajería con queues por servicio'],
    ['Kali-Auditor', '2', '4', '40', 'Herramientas pesadas de auditoría (Burp Suite, Metasploit)'],
    ['Metasploit-Edu', '1', '2', '20', 'Servidor vulnerable de práctica'],
    ['Total', '13', '25', '268', ''],
]
add_table(res_headers, res_rows, col_widths=[2.8, 1.8, 1.5, 1.8, 8.9])

# --- Security Tools ---
add_heading_styled('Herramientas de validación de seguridad', level=2)

sec_headers = ['Herramienta', 'VM', 'Propósito']
sec_rows = [
    ['Nmap', 'Kali-Auditor', 'Escaneo de puertos y detección de servicios expuestos'],
    ['Nikto', 'Kali-Auditor', 'Análisis de vulnerabilidades en el servidor web y la API'],
    ['Burp Suite', 'Kali-Auditor', 'Interceptación y manipulación de requests JWT, validación de autorización por roles'],
    ['SQLMap', 'Kali-Auditor', 'Pruebas de inyección SQL sobre endpoints de la API'],
    ['OWASP ZAP', 'Kali-Auditor', 'Escaneo automatizado de vulnerabilidades OWASP Top 10'],
    ['Metasploit Framework', 'Kali-Auditor', 'Explotación controlada contra Metasploit-Edu para validar segmentación de red'],
]
add_table(sec_headers, sec_rows, col_widths=[3.5, 2.5, 10.8])

# --- Save ---
output_path = r'c:\Users\EQUIPO\Downloads\proyecto\barbershop\docs\Configuracion_Entorno_Virtual.docx'
doc.save(output_path)
print(f'Documento generado: {output_path}')
